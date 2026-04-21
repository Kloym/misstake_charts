import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

TARGET_URL = "https://swimparser.ru/" 
OUTPUT_FILENAME = "Ссылка_на_сайт_Swimparser.docx"

def create_bulletproof_word_link(url, filename):
    doc = docx.Document()

    for _ in range(8):
        doc.add_paragraph()

    p_instruction = doc.add_paragraph()
    p_instruction.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_instruction.add_run("ВНИМАНИЕ: Если ссылка не нажимается, нажмите желтую кнопку «Редактировать копию» сверху или скачайте файл. Также вы можете просто скопировать эту ссылку")
    run_inst.font.color.rgb = docx.shared.RGBColor(150, 150, 150)
    
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ПЕРЕЙТИ НА САЙТ SWIMPARSER:")
    run_title.bold = True
    run_title.font.size = docx.shared.Pt(18)

    p_link = doc.add_paragraph()
    p_link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_link = p_link.add_run(url)
    run_link.font.color.rgb = docx.shared.RGBColor(0, 112, 255)
    run_link.underline = True
    run_link.font.size = docx.shared.Pt(22)
    
    doc.save(filename)
    print(f"✅ Обновленный файл '{filename}' создан!")

if __name__ == "__main__":
    create_bulletproof_word_link(TARGET_URL, OUTPUT_FILENAME)